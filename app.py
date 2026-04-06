import io
import json
import os
import secrets
from datetime import datetime
from typing import Any, Dict, Optional

from docx import Document
from flask import Flask, jsonify, render_template, request, send_file, send_from_directory, session
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

from config import BASE_DIR, Config
from models import EntitlementGrant, ExportLog, Subscription, TrainingFeedback, TrainingPlan, User, UserProfile, WebhookEvent, db, now_str
from rules import analyze_feedback, build_rollover_profile, compute_css_from_inputs, generate_plan, growth_stage_from_progress, normalize_text
from storage import StorageService

try:  # pragma: no cover
    import stripe
except Exception:  # pragma: no cover
    stripe = None


STATIC_DIR = BASE_DIR / 'static'
TEMPLATE_DIR = BASE_DIR / 'templates'


app = Flask(__name__, static_folder=str(STATIC_DIR), template_folder=str(TEMPLATE_DIR))
app.config.from_object(Config)
db.init_app(app)
storage = StorageService(app)


# ---------- helpers ----------

def json_load(value: Optional[str]):
    if not value:
        return None
    try:
        return json.loads(value)
    except Exception:
        return None


def json_dump(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False)


def model_to_profile_dict(profile: UserProfile) -> Dict[str, Any]:
    return {
        'id': profile.id,
        'user_id': profile.user_id,
        'created_at': profile.created_at,
        'user_group': profile.user_group,
        'goal': profile.goal,
        'plan_scope': profile.plan_scope,
        'cycle_length': profile.cycle_length,
        'display_name': profile.display_name,
        'sex': profile.sex,
        'age': profile.age,
        'height_cm': profile.height_cm,
        'weight_kg': profile.weight_kg,
        'body_fat': profile.body_fat,
        'activity_level': profile.activity_level,
        'level': profile.level,
        'primary_stroke': profile.primary_stroke,
        'event_focus': profile.event_focus,
        'sessions_per_week': profile.sessions_per_week,
        'session_duration': profile.session_duration,
        'discomfort': profile.discomfort,
        'recent_result_flag': profile.recent_result_flag,
        'best_result_distance': profile.best_result_distance,
        'best_result_time': profile.best_result_time,
        'css_400': profile.css_400,
        'css_200': profile.css_200,
        'css_pace_seconds': profile.css_pace_seconds,
        'watch_data_flag': profile.watch_data_flag,
        'extra_note': profile.extra_note,
    }


def plan_model_to_dict(plan: TrainingPlan) -> Dict[str, Any]:
    return {
        'id': plan.id,
        'user_id': plan.user_id,
        'profile_id': plan.profile_id,
        'created_at': plan.created_at,
        'plan_scope': plan.plan_scope,
        'cycle_length': plan.cycle_length,
        'plan_version': plan.plan_version,
        'status': plan.status,
        'summary': plan.summary,
        'data_json': json_load(plan.data_json) or {},
        'calories_json': json_load(plan.calories_json),
        'nutrition_json': json_load(plan.nutrition_json),
    }


def current_user() -> User:
    uid = session.get('user_id')
    user = db.session.get(User, uid) if uid else None
    if user:
        return user

    guest = User(
        nickname=f'访客-{secrets.token_hex(2)}',
        is_guest=True,
        membership_tier='free',
        created_at=now_str(),
        updated_at=now_str(),
    )
    db.session.add(guest)
    db.session.commit()
    session['user_id'] = guest.id
    return guest


def require_user() -> User:
    return current_user()


def parse_dt(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    try:
        return datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
    except Exception:
        return None


BASE_FEATURES = {
    'free': {
        'basic_plan': True,
        'cycle_plan': False,
        'exports': bool(app.config['ALLOW_FREE_EXPORTS']),
        'progress_dashboard': False,
        'progress_trends_full': False,
        'watch_feedback_enhanced': False,
        'nutrition_advanced': False,
        'rollover_week': False,
    },
    'vip': {
        'basic_plan': True,
        'cycle_plan': True,
        'exports': True,
        'progress_dashboard': True,
        'progress_trends_full': True,
        'watch_feedback_enhanced': True,
        'nutrition_advanced': True,
        'rollover_week': True,
    }
}


def compute_entitlements(user: User) -> Dict[str, bool]:
    tier = user.membership_tier or 'free'
    features = dict(BASE_FEATURES.get(tier, BASE_FEATURES['free']))
    now = datetime.utcnow()
    grants = EntitlementGrant.query.filter_by(user_id=user.id, is_enabled=True).all()
    for grant in grants:
        expires = parse_dt(grant.expires_at)
        if expires and expires < now:
            continue
        features[grant.feature_code] = True
    return features


def has_feature(user: User, feature_code: str) -> bool:
    return bool(compute_entitlements(user).get(feature_code))


def is_vip(user: User) -> bool:
    return (user.membership_tier or 'free') == 'vip'


def export_allowed(user: User) -> bool:
    return has_feature(user, 'exports')


def latest_profile(user_id: int) -> Optional[UserProfile]:
    return UserProfile.query.filter_by(user_id=user_id).order_by(UserProfile.id.desc()).first()


def fetch_profile(profile_id: int) -> Optional[UserProfile]:
    return db.session.get(UserProfile, profile_id)


def fetch_plan(plan_id: int) -> Optional[TrainingPlan]:
    return db.session.get(TrainingPlan, plan_id)


def current_plan_model(user_id: int) -> Optional[TrainingPlan]:
    return TrainingPlan.query.filter_by(user_id=user_id, status='active').order_by(TrainingPlan.id.desc()).first()


def parse_completion_ratio(status: str) -> float:
    return {
        '完成了全部训练': 1.0,
        '完成了一大部分': 0.75,
        '只完成了一小部分': 0.35,
        '没有完成': 0.0,
    }.get(status, 0.6)


def serialize_profile_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    cleaned = {
        'user_group': normalize_text(data.get('user_group')) or '爱好者',
        'goal': normalize_text(data.get('goal')) or '提高耐力',
        'plan_scope': normalize_text(data.get('plan_scope')) or 'week',
        'cycle_length': int(data.get('cycle_length') or 4),
        'display_name': normalize_text(data.get('display_name')),
        'sex': normalize_text(data.get('sex')),
        'age': int(data['age']) if data.get('age') not in [None, ''] else None,
        'height_cm': float(data['height_cm']) if data.get('height_cm') not in [None, ''] else None,
        'weight_kg': float(data['weight_kg']) if data.get('weight_kg') not in [None, ''] else None,
        'body_fat': float(data['body_fat']) if data.get('body_fat') not in [None, ''] else None,
        'activity_level': normalize_text(data.get('activity_level')) or '轻度活动',
        'level': normalize_text(data.get('level')) or '爱好者',
        'primary_stroke': normalize_text(data.get('primary_stroke')),
        'event_focus': normalize_text(data.get('event_focus')),
        'sessions_per_week': int(data.get('sessions_per_week') or 3),
        'session_duration': int(data.get('session_duration') or 45),
        'discomfort': normalize_text(data.get('discomfort')) or '没有明显不适',
        'recent_result_flag': normalize_text(data.get('recent_result_flag')) or '无',
        'best_result_distance': normalize_text(data.get('best_result_distance')),
        'best_result_time': normalize_text(data.get('best_result_time')),
        'css_400': normalize_text(data.get('css_400')),
        'css_200': normalize_text(data.get('css_200')),
        'watch_data_flag': normalize_text(data.get('watch_data_flag')) or '无',
        'extra_note': normalize_text(data.get('extra_note')),
    }
    css_info = compute_css_from_inputs(cleaned['css_400'], cleaned['css_200'])
    cleaned['css_info'] = css_info
    cleaned['css_pace_seconds'] = css_info.get('css_pace_seconds') if css_info.get('valid') else None
    return cleaned


def create_profile_model(user_id: int, payload: Dict[str, Any]) -> UserProfile:
    profile = UserProfile(
        user_id=user_id,
        created_at=now_str(),
        user_group=payload['user_group'],
        goal=payload['goal'],
        plan_scope=payload['plan_scope'],
        cycle_length=payload['cycle_length'],
        display_name=payload['display_name'],
        sex=payload['sex'],
        age=payload['age'],
        height_cm=payload['height_cm'],
        weight_kg=payload['weight_kg'],
        body_fat=payload['body_fat'],
        activity_level=payload['activity_level'],
        level=payload['level'],
        primary_stroke=payload['primary_stroke'],
        event_focus=payload['event_focus'],
        sessions_per_week=payload['sessions_per_week'],
        session_duration=payload['session_duration'],
        discomfort=payload['discomfort'],
        recent_result_flag=payload['recent_result_flag'],
        best_result_distance=payload['best_result_distance'],
        best_result_time=payload['best_result_time'],
        css_400=payload['css_400'],
        css_200=payload['css_200'],
        css_pace_seconds=payload['css_pace_seconds'],
        watch_data_flag=payload['watch_data_flag'],
        extra_note=payload['extra_note'],
    )
    db.session.add(profile)
    db.session.commit()
    return profile


def save_plan(user_id: int, profile: UserProfile, plan_data: Dict[str, Any]) -> TrainingPlan:
    current_active = TrainingPlan.query.filter_by(user_id=user_id, status='active').all()
    for p in current_active:
        p.status = 'archived'
    latest_version = db.session.query(db.func.max(TrainingPlan.plan_version)).filter_by(user_id=user_id).scalar() or 0
    plan = TrainingPlan(
        user_id=user_id,
        profile_id=profile.id,
        created_at=now_str(),
        plan_scope=profile.plan_scope,
        cycle_length=profile.cycle_length,
        plan_version=int(latest_version) + 1,
        status='active',
        summary=plan_data['summary'],
        data_json=json_dump(plan_data),
        calories_json=json_dump(plan_data.get('calories')) if plan_data.get('calories') else None,
        nutrition_json=json_dump(plan_data.get('nutrition')) if plan_data.get('nutrition') else None,
    )
    db.session.add(plan)
    db.session.commit()
    return plan


def build_docx_bytes(plan_data: Dict[str, Any]) -> io.BytesIO:
    doc = Document()
    doc.add_heading('SwimDiary 训练计划', level=0)
    doc.add_paragraph(plan_data['summary'])
    doc.add_paragraph(plan_data['headline'])
    doc.add_heading('本周计划', level=1)
    for session_data in plan_data['weekly_plan']:
        doc.add_heading(session_data['title'], level=2)
        doc.add_paragraph(session_data['intention'])
        doc.add_paragraph(f"预计总量：{session_data['total_distance']}")
        doc.add_paragraph(f"执行说明：{session_data['effort_explainer']}")
        for segment in session_data['segments']:
            doc.add_paragraph(f"{segment['label']}｜{segment['zone']}：{segment['detail']}", style='List Bullet')
        for note in session_data['notes']:
            doc.add_paragraph(note, style='List Bullet 2')
    if plan_data.get('cycle_overview'):
        doc.add_heading('周期总览', level=1)
        for week in plan_data['cycle_overview']:
            doc.add_paragraph(f"{week['title']}｜{week['theme']}｜{week['focus']}", style='List Bullet')
    doc.add_heading('强度说明', level=1)
    for item in plan_data['intensity_guide']:
        doc.add_paragraph(f"{item['code']}｜{item['name']}｜{item['pace_range']}｜{item['description']}", style='List Bullet')
    if plan_data.get('nutrition'):
        doc.add_heading('营养与减脂建议', level=1)
        doc.add_paragraph(plan_data['nutrition']['summary'])
        for meal in plan_data['nutrition']['meals']:
            doc.add_paragraph(meal, style='List Bullet')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf_bytes(plan_data: Dict[str, Any]) -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title='SwimDiary 训练计划')
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CNTitle', parent=styles['Title'], fontName='STSong-Light', leading=24)
    h1_style = ParagraphStyle('CNH1', parent=styles['Heading1'], fontName='STSong-Light', leading=18)
    h2_style = ParagraphStyle('CNH2', parent=styles['Heading2'], fontName='STSong-Light', leading=16)
    body_style = ParagraphStyle('CNBody', parent=styles['BodyText'], fontName='STSong-Light', leading=14)
    story = [
        Paragraph('SwimDiary 训练计划', title_style), Spacer(1, 10),
        Paragraph(plan_data['summary'], body_style), Spacer(1, 8),
        Paragraph(plan_data['headline'], body_style), Spacer(1, 12),
        Paragraph('本周计划', h1_style)
    ]
    for session_data in plan_data['weekly_plan']:
        story += [
            Paragraph(session_data['title'], h2_style),
            Paragraph(session_data['intention'], body_style),
            Paragraph(f"预计总量：{session_data['total_distance']}", body_style)
        ]
        for segment in session_data['segments']:
            story.append(Paragraph(f"{segment['label']}｜{segment['zone']}：{segment['detail']}", body_style))
        story.append(Spacer(1, 6))
    if plan_data.get('cycle_overview'):
        story.append(Paragraph('周期总览', h1_style))
        for week in plan_data['cycle_overview']:
            story.append(Paragraph(f"{week['title']}｜{week['theme']}｜{week['focus']}", body_style))
    if plan_data.get('nutrition'):
        story.append(Paragraph('营养与减脂建议', h1_style))
        story.append(Paragraph(plan_data['nutrition']['summary'], body_style))
    doc.build(story)
    buf.seek(0)
    return buf


def log_export(user_id: int, plan_id: int, export_type: str, file_name: str, storage_key: Optional[str] = None, storage_url: Optional[str] = None):
    db.session.add(ExportLog(
        user_id=user_id,
        plan_id=plan_id,
        created_at=now_str(),
        export_type=export_type,
        file_name=file_name,
        storage_key=storage_key,
        storage_url=storage_url,
    ))
    db.session.commit()


# ---------- billing ----------

def stripe_ready() -> bool:
    return stripe is not None and bool(app.config['STRIPE_SECRET_KEY'])


def configure_stripe():
    if stripe_ready():
        stripe.api_key = app.config['STRIPE_SECRET_KEY']


def unix_to_str(ts: Optional[int]) -> Optional[str]:
    if not ts:
        return None
    return datetime.utcfromtimestamp(int(ts)).strftime('%Y-%m-%d %H:%M:%S')


def refresh_membership_from_subscriptions(user: User):
    active_status = {'active', 'trialing', 'past_due'}
    has_vip = Subscription.query.filter(
        Subscription.user_id == user.id,
        Subscription.status.in_(active_status)
    ).count() > 0
    user.membership_tier = 'vip' if has_vip else 'free'
    user.updated_at = now_str()
    db.session.commit()


def upsert_subscription_for_user(user: User, stripe_subscription_obj: Dict[str, Any]):
    sub_id = stripe_subscription_obj['id']
    sub = Subscription.query.filter_by(stripe_subscription_id=sub_id).first()
    if not sub:
        sub = Subscription(user_id=user.id, provider='stripe', created_at=now_str())
        db.session.add(sub)
    sub.user_id = user.id
    sub.provider = 'stripe'
    sub.stripe_customer_id = stripe_subscription_obj.get('customer')
    sub.stripe_subscription_id = sub_id
    items = stripe_subscription_obj.get('items', {}).get('data', [])
    sub.stripe_price_id = items[0]['price']['id'] if items else None
    sub.status = stripe_subscription_obj.get('status') or 'inactive'
    sub.current_period_end = unix_to_str(stripe_subscription_obj.get('current_period_end'))
    sub.cancel_at_period_end = bool(stripe_subscription_obj.get('cancel_at_period_end'))
    sub.updated_at = now_str()
    if stripe_subscription_obj.get('customer'):
        user.stripe_customer_id = stripe_subscription_obj.get('customer')
    refresh_membership_from_subscriptions(user)


def process_stripe_event(event: Dict[str, Any]):
    event_id = event['id']
    if WebhookEvent.query.filter_by(provider='stripe', event_id=event_id).first():
        return
    event_type = event['type']
    obj = event['data']['object']

    db.session.add(WebhookEvent(provider='stripe', event_id=event_id, event_type=event_type, processed_at=now_str()))
    db.session.commit()

    if event_type == 'checkout.session.completed':
        user_id = obj.get('metadata', {}).get('user_id')
        if user_id:
            user = db.session.get(User, int(user_id))
            if user and obj.get('customer'):
                user.stripe_customer_id = obj.get('customer')
                user.updated_at = now_str()
                db.session.commit()
            subscription_id = obj.get('subscription')
            if subscription_id and stripe_ready() and user:
                configure_stripe()
                sub_obj = stripe.Subscription.retrieve(subscription_id)
                upsert_subscription_for_user(user, sub_obj)
        return

    if event_type.startswith('customer.subscription.'):
        customer_id = obj.get('customer')
        user = User.query.filter_by(stripe_customer_id=customer_id).first()
        if user:
            upsert_subscription_for_user(user, obj)
        return

    if event_type == 'invoice.payment_failed':
        customer_id = obj.get('customer')
        user = User.query.filter_by(stripe_customer_id=customer_id).first()
        if not user:
            return
        sub_id = obj.get('subscription')
        sub = Subscription.query.filter_by(stripe_subscription_id=sub_id).first()
        if sub:
            sub.status = 'past_due'
            sub.updated_at = now_str()
            db.session.commit()
        refresh_membership_from_subscriptions(user)
        return

    if event_type == 'invoice.paid':
        customer_id = obj.get('customer')
        user = User.query.filter_by(stripe_customer_id=customer_id).first()
        if not user or not stripe_ready():
            return
        sub_id = obj.get('subscription')
        if not sub_id:
            return
        configure_stripe()
        sub_obj = stripe.Subscription.retrieve(sub_id)
        upsert_subscription_for_user(user, sub_obj)


# ---------- analytics / progress ----------

def summarize_progress(user_id: int) -> Dict[str, Any]:
    user = db.session.get(User, user_id)
    entitlements = compute_entitlements(user) if user else dict(BASE_FEATURES['free'])
    plan_rows = TrainingPlan.query.filter_by(user_id=user_id).order_by(TrainingPlan.id.desc()).limit(20).all()
    feedback_rows = TrainingFeedback.query.filter_by(user_id=user_id).order_by(TrainingFeedback.id.desc()).limit(40).all()
    export_rows = ExportLog.query.filter_by(user_id=user_id).order_by(ExportLog.id.desc()).limit(12).all()

    latest_profile_model = latest_profile(user_id)
    latest_profile_data = model_to_profile_dict(latest_profile_model) if latest_profile_model else None

    plan_history = []
    latest_css = None
    latest_plan_scope = None
    for row in plan_rows:
        data = json_load(row.data_json) or {}
        plan_history.append({
            'plan_id': row.id,
            'created_at': row.created_at,
            'plan_scope': row.plan_scope,
            'plan_version': row.plan_version,
            'summary': row.summary,
            'headline': data.get('headline'),
        })
        latest_plan_scope = latest_plan_scope or row.plan_scope
        for item in data.get('intensity_guide', []):
            if item.get('code') == 'EN2':
                latest_css = item.get('pace_range')
                break

    timeline = []
    completion_rates, fatigue_vals, distances, calorie_gaps = [], [], [], []
    sleep_good = sleep_poor = risk_count = 0
    current_streak = best_streak = 0
    adherence_score = 0

    for row in reversed(feedback_rows):
        result = json_load(row.result_json) or {}
        metrics = json_load(row.screenshot_metrics_json) or {}
        ratio = row.completion_ratio if row.completion_ratio is not None else parse_completion_ratio(row.completion_status)
        ratio = max(0.0, min(float(ratio), 1.0))
        completion_rates.append(float(ratio))
        fatigue_vals.append(int(row.fatigue_score))
        distances.append(int(metrics.get('actual_distance_m') or 0))
        if row.sleep_status in ['7 小时以上', '6–7 小时']:
            sleep_good += 1
        else:
            sleep_poor += 1
        if result.get('next_week_action') in ['回调', '维持或小幅回调']:
            risk_count += 1
        if ratio >= 0.7:
            current_streak += 1
            best_streak = max(best_streak, current_streak)
        else:
            current_streak = 0
        adherence_score += ratio * 100
        timeline.append({
            'feedback_id': row.id,
            'created_at': row.created_at,
            'session_index': row.session_index,
            'completion_status': row.completion_status,
            'completion_ratio': round(float(ratio) * 100),
            'fatigue_score': row.fatigue_score,
            'sleep_status': row.sleep_status,
            'discomfort_feedback': row.discomfort_feedback,
            'judgment': result.get('judgment'),
            'direction': result.get('direction'),
            'next_week_action': result.get('next_week_action'),
            'watch_distance': metrics.get('actual_distance_m'),
            'watch_analysis': result.get('watch_analysis'),
        })

    avg_completion = round(sum(completion_rates) / len(completion_rates) * 100) if completion_rates else 0
    avg_fatigue = round(sum(fatigue_vals) / len(fatigue_vals), 1) if fatigue_vals else 0
    recovery_score = max(0, min(100, round((100 - (avg_fatigue * 9)) + (sleep_good - sleep_poor) * 3))) if fatigue_vals else 0
    consistency_score = max(0, min(100, round(avg_completion * 0.75 + min(best_streak, 6) * 4))) if completion_rates else 0

    recent_completion = sum(completion_rates[-3:]) / len(completion_rates[-3:]) * 100 if completion_rates[-3:] else avg_completion
    previous_completion = sum(completion_rates[-6:-3]) / len(completion_rates[-6:-3]) * 100 if completion_rates[-6:-3] else recent_completion
    recent_fatigue = sum(fatigue_vals[-3:]) / len(fatigue_vals[-3:]) if fatigue_vals[-3:] else avg_fatigue
    previous_fatigue = sum(fatigue_vals[-6:-3]) / len(fatigue_vals[-6:-3]) if fatigue_vals[-6:-3] else recent_fatigue

    if risk_count >= 2 or avg_fatigue >= 7.5:
        status = '需要回调'
    elif avg_completion >= 80 and avg_fatigue <= 5.5:
        status = '可以推进'
    else:
        status = '维持推进'

    if recent_completion - previous_completion >= 8 and recent_fatigue <= previous_fatigue + 0.5:
        momentum = '上升中'
    elif previous_completion - recent_completion >= 8 or recent_fatigue - previous_fatigue >= 1.2:
        momentum = '需要减压'
    else:
        momentum = '平稳积累'

    growth_score = max(0, min(100, round(avg_completion * 0.42 + recovery_score * 0.28 + consistency_score * 0.3)))
    stage = growth_stage_from_progress({'overview': {'feedback_count': len(timeline), 'avg_completion_rate': avg_completion, 'avg_fatigue': avg_fatigue, 'current_status': status}})

    milestones = []
    if len(plan_history) >= 1:
        milestones.append('已建立专属训练档案')
    if len(timeline) >= 3:
        milestones.append('已形成初步反馈闭环')
    if avg_completion >= 80 and len(timeline) >= 3:
        milestones.append('训练执行率进入稳定区间')
    if latest_profile_data and latest_profile_data.get('css_pace_seconds'):
        milestones.append('已接入 CSS 强度参考')
    if len(export_rows) >= 1:
        milestones.append('已开始导出并带着计划执行')

    # Nutrition / fat-loss glance from latest active plan
    calorie_overview = None
    active_plan = current_plan_model(user_id)
    if active_plan:
        data = plan_model_to_dict(active_plan)
        calories = data.get('calories_json') or data['data_json'].get('calories')
        if calories and entitlements.get('nutrition_advanced'):
            calorie_overview = calories
            if calories.get('target_deficit_kcal'):
                gap = int(calories.get('target_deficit_kcal'))
                calorie_gaps = [gap for _ in completion_rates[-6:]]

    visible_feedback = timeline[-10:] if entitlements.get('progress_trends_full') else timeline[-4:]
    visible_plans = plan_history[:10] if entitlements.get('progress_trends_full') else plan_history[:3]

    next_focus = {
        '需要回调': '下一周优先恢复节奏与完成率，避免额外增加刺激。',
        '可以推进': '下一周可以把主训练稍微做满，并保留一节关键推进课。',
        '维持推进': '下一周继续稳住当前节奏，在保证恢复的前提下小幅优化质量。',
    }[status]

    return {
        'plan_history': visible_plans,
        'feedback_history': list(reversed(visible_feedback)),
        'export_history': [{
            'id': r.id, 'created_at': r.created_at, 'export_type': r.export_type,
            'file_name': r.file_name, 'storage_url': r.storage_url
        } for r in export_rows],
        'overview': {
            'plans_count': len(plan_history),
            'feedback_count': len(timeline),
            'avg_completion_rate': avg_completion,
            'avg_fatigue': avg_fatigue,
            'sleep_good_count': sleep_good,
            'sleep_poor_count': sleep_poor,
            'current_status': status,
            'momentum_label': momentum,
            'growth_score': growth_score,
            'consistency_score': consistency_score,
            'recovery_score': recovery_score,
            'best_streak': best_streak,
            'latest_css_hint': latest_css,
            'stage': stage,
            'next_focus': next_focus,
            'rollover_ready': len(timeline) >= 2,
            'latest_plan_scope': latest_plan_scope or 'week',
        },
        'charts': {
            'labels': [item['created_at'][5:10] for item in timeline[-8:]],
            'completion': [item['completion_ratio'] for item in timeline[-8:]],
            'fatigue': [item['fatigue_score'] for item in timeline[-8:]],
            'distance': [int(item['watch_distance'] or 0) for item in timeline[-8:]],
            'calorie_gap': calorie_gaps[-8:] if calorie_gaps else [],
        },
        'milestones': milestones[:5],
        'calorie_overview': calorie_overview if entitlements.get('nutrition_advanced') else None,
        'entitlements': entitlements,
    }


# ---------- routes ----------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/manifest.json')
def manifest():
    return send_from_directory(STATIC_DIR, 'manifest.json')


@app.route('/service-worker.js')
def sw():
    return send_from_directory(STATIC_DIR, 'service-worker.js')


@app.route('/api/health')
def health():
    return jsonify({'ok': True, 'app': 'SwimDiary Production Ready', 'time': now_str()})


@app.route('/api/options')
def options():
    return jsonify({
        'user_groups': ['爱好者', '进阶训练者', '备赛用户', '减脂用户'],
        'goals': ['提高耐力', '提高速度', '技术改善', '备赛', '减脂'],
        'levels': ['零基础/恢复训练', '爱好者', '进阶爱好者', '比赛/备赛'],
        'plan_scopes': [
            {'value': 'week', 'label': '默认：本周计划'},
            {'value': 'cycle', 'label': '进阶：4/6/8 周周期计划'}
        ],
        'cycle_lengths': [4, 6, 8],
        'activity_levels': ['久坐', '轻度活动', '中度活动', '较高活动'],
        'sleep_status': ['5 小时以下', '5–6 小时', '6–7 小时', '7 小时以上'],
        'completion_status': ['完成了全部训练', '完成了一大部分', '只完成了一小部分', '没有完成'],
        'discomfort_feedback': ['没有明显不适', '轻微不适', '明显不适'],
        'pricing': {
            'free': ['体验模式使用', '基础建档', '基础周计划', '基础反馈', '简版成长看板'],
            'vip': ['滚动周计划', '周期计划', 'Word/PDF 导出', '完整成长看板', '减脂联动', '截图反馈增强']
        }
    })


@app.route('/api/me')
def me():
    user = current_user()
    profile = latest_profile(user.id)
    current = current_plan_model(user.id)
    return jsonify({
        'user': {
            'id': user.id,
            'email': user.email,
            'nickname': user.nickname,
            'is_guest': bool(user.is_guest),
            'membership_tier': user.membership_tier or 'free',
        },
        'has_profile': profile is not None,
        'has_plan': current is not None,
        'latest_profile_id': profile.id if profile else None,
        'current_plan_id': current.id if current else None,
    })


@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    user = current_user()
    data = request.get_json(force=True)
    email = normalize_text(data.get('email')).lower()
    password = data.get('password')
    nickname = normalize_text(data.get('nickname')) or user.nickname
    if not email or not password:
        return jsonify({'error': '请填写邮箱和密码。'}), 400
    existed = User.query.filter(User.email == email, User.id != user.id).first()
    if existed:
        return jsonify({'error': '该邮箱已被注册。'}), 400
    user.email = email
    user.password_hash = generate_password_hash(password)
    user.nickname = nickname
    user.is_guest = False
    user.updated_at = now_str()
    db.session.commit()
    return jsonify({'ok': True, 'message': '账号已创建，当前历史数据已绑定。'})


@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    data = request.get_json(force=True)
    email = normalize_text(data.get('email')).lower()
    password = data.get('password') or ''
    user = User.query.filter_by(email=email).first()
    if not user or not user.password_hash or not check_password_hash(user.password_hash, password):
        return jsonify({'error': '邮箱或密码不正确。'}), 400
    session['user_id'] = user.id
    return jsonify({'ok': True, 'message': '登录成功。'})


@app.route('/api/auth/logout', methods=['POST'])
def auth_logout():
    session.clear()
    current_user()
    return jsonify({'ok': True, 'message': '已退出当前账号，已切换为访客模式。'})


@app.route('/api/css/compute', methods=['POST'])
def css_compute():
    data = request.get_json(force=True)
    return jsonify(compute_css_from_inputs(data.get('css_400'), data.get('css_200')))


@app.route('/api/profile', methods=['POST'])
def create_profile():
    user = require_user()
    data = request.get_json(force=True)
    profile_payload = serialize_profile_payload(data)
    required = ['user_group', 'goal', 'level', 'sessions_per_week', 'session_duration']
    for key in required:
        if profile_payload.get(key) in [None, '', 0]:
            return jsonify({'error': f'缺少字段: {key}'}), 400
    profile = create_profile_model(user.id, profile_payload)
    return jsonify({'profile_id': profile.id, 'css_info': profile_payload['css_info']})


@app.route('/api/profile/latest')
def profile_latest():
    user = require_user()
    profile = latest_profile(user.id)
    return jsonify({'profile': model_to_profile_dict(profile) if profile else None})


@app.route('/api/plan/generate', methods=['POST'])
def plan_generate():
    user = require_user()
    data = request.get_json(force=True)
    profile_id = int(data.get('profile_id') or 0)
    profile = fetch_profile(profile_id) if profile_id else latest_profile(user.id)
    if not profile or profile.user_id != user.id:
        return jsonify({'error': '用户建档不存在'}), 404
    if profile.plan_scope == 'cycle' and not has_feature(user, 'cycle_plan'):
        return jsonify({'error': '周期计划与更深度的成长方案为 VIP 功能。你仍可以先体验默认周计划。'}), 403
    plan_data = generate_plan(model_to_profile_dict(profile))
    plan = save_plan(user.id, profile, plan_data)
    return jsonify({'plan_id': plan.id, **plan_data})


@app.route('/api/plan/current')
def plan_current():
    user = require_user()
    plan = current_plan_model(user.id)
    if not plan:
        return jsonify({'plan': None})
    data = plan_model_to_dict(plan)
    return jsonify({'plan_id': plan.id, **data['data_json']})


@app.route('/api/plan/rollover', methods=['POST'])
def plan_rollover():
    user = require_user()
    if not has_feature(user, 'rollover_week'):
        return jsonify({'error': '滚动周计划为 VIP 功能。'}), 403
    profile = latest_profile(user.id)
    if not profile:
        return jsonify({'error': '请先完成建档。'}), 400
    progress = summarize_progress(user.id)
    current_plan = current_plan_model(user.id)
    if not current_plan:
        return jsonify({'error': '请先生成当前周计划。'}), 400

    rollover_payload = build_rollover_profile(model_to_profile_dict(profile), progress)
    rollover_payload['plan_scope'] = 'week'
    css_info = compute_css_from_inputs(rollover_payload.get('css_400'), rollover_payload.get('css_200'))
    rollover_payload['css_info'] = css_info
    rollover_payload['css_pace_seconds'] = css_info.get('css_pace_seconds') if css_info.get('valid') else rollover_payload.get('css_pace_seconds')

    new_profile = create_profile_model(user.id, rollover_payload)
    plan_data = generate_plan(model_to_profile_dict(new_profile))
    plan_data['headline'] = '这是根据你最近的反馈自动滚动生成的下一周计划，重点不是推翻一切，而是在可恢复的前提下持续推进。'
    plan = save_plan(user.id, new_profile, plan_data)
    return jsonify({
        'message': '已根据最近反馈生成下一周滚动计划。',
        'plan_id': plan.id,
        'rollover_note': rollover_payload.get('extra_note'),
        **plan_data
    })


@app.route('/api/plan/history')
def plan_history():
    user = require_user()
    return jsonify(summarize_progress(user.id))


@app.route('/api/plan/<int:plan_id>')
def plan_detail(plan_id: int):
    user = require_user()
    plan = fetch_plan(plan_id)
    if not plan or plan.user_id != user.id:
        return jsonify({'error': '计划不存在'}), 404
    data = plan_model_to_dict(plan)
    return jsonify({'plan_id': plan_id, **data['data_json']})


@app.route('/api/progress/overview')
def progress_overview():
    user = require_user()
    return jsonify(summarize_progress(user.id))


@app.route('/api/export/docx/<int:plan_id>')
def export_docx(plan_id: int):
    user = require_user()
    if not export_allowed(user):
        return jsonify({'error': 'Word 导出为 VIP 功能。'}), 403
    plan = fetch_plan(plan_id)
    if not plan or plan.user_id != user.id:
        return jsonify({'error': '计划不存在'}), 404
    data = plan_model_to_dict(plan)['data_json']
    file_name = f'SwimDiary_plan_{plan_id}.docx'
    buf = build_docx_bytes(data)
    storage_key, storage_url = storage.upload_bytes(buf.getvalue(), key=f'exports/user_{user.id}/{file_name}', filename=file_name,
                                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    buf.seek(0)
    log_export(user.id, plan_id, 'docx', file_name, storage_key=storage_key, storage_url=storage_url)
    return send_file(buf, as_attachment=True, download_name=file_name,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/export/pdf/<int:plan_id>')
def export_pdf(plan_id: int):
    user = require_user()
    if not export_allowed(user):
        return jsonify({'error': 'PDF 导出为 VIP 功能。'}), 403
    plan = fetch_plan(plan_id)
    if not plan or plan.user_id != user.id:
        return jsonify({'error': '计划不存在'}), 404
    data = plan_model_to_dict(plan)['data_json']
    file_name = f'SwimDiary_plan_{plan_id}.pdf'
    buf = build_pdf_bytes(data)
    storage_key, storage_url = storage.upload_bytes(buf.getvalue(), key=f'exports/user_{user.id}/{file_name}', filename=file_name,
                                                    content_type='application/pdf')
    buf.seek(0)
    log_export(user.id, plan_id, 'pdf', file_name, storage_key=storage_key, storage_url=storage_url)
    return send_file(buf, as_attachment=True, download_name=file_name, mimetype='application/pdf')


@app.route('/api/feedback', methods=['POST'])
def feedback():
    user = require_user()
    data = request.form.to_dict() if request.form else (request.get_json(force=True) if request.is_json else {})
    plan_id = int(data.get('plan_id') or 0)
    session_index = int(data.get('session_index') or 0)
    if not plan_id:
        return jsonify({'error': '缺少 plan_id。'}), 400
    plan = fetch_plan(plan_id)
    if not plan or plan.user_id != user.id:
        return jsonify({'error': '相关计划不存在。'}), 404
    profile = fetch_profile(plan.profile_id)
    if not profile:
        return jsonify({'error': '相关建档不存在。'}), 404

    screenshot_key = None
    screenshot_url = None
    if 'screenshot' in request.files and request.files['screenshot'].filename:
        file = request.files['screenshot']
        safe_name = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{secure_filename(file.filename)}"
        screenshot_key, screenshot_url = storage.upload_file_storage(file, key=f'watch-feedback/user_{user.id}/{safe_name}')

    metrics = {
        'actual_distance_m': int(data['actual_distance_m']) if data.get('actual_distance_m') else None,
        'actual_duration_s': int(data['actual_duration_s']) if data.get('actual_duration_s') else None,
        'avg_pace_per_100_s': float(data['avg_pace_per_100_s']) if data.get('avg_pace_per_100_s') else None,
        'avg_hr': int(data['avg_hr']) if data.get('avg_hr') else None,
        'ocr_text': normalize_text(data.get('ocr_text')),
    }
    payload = {
        'session_index': session_index,
        'completion_status': data.get('completion_status') or '完成了一大部分',
        'fatigue_score': int(data.get('fatigue_score') or 5),
        'discomfort_feedback': data.get('discomfort_feedback') or '没有明显不适',
        'sleep_status': data.get('sleep_status') or '6–7 小时',
        'feedback_note': normalize_text(data.get('feedback_note')),
        'screenshot_metrics': metrics,
    }
    result = analyze_feedback(model_to_profile_dict(profile), plan_model_to_dict(plan)['data_json'], payload)
    ratio = result.get('watch_analysis', {}).get('completion_ratio')
    if ratio is None:
        ratio = parse_completion_ratio(payload['completion_status'])
    ratio = max(0.0, min(float(ratio), 1.0))

    feedback_row = TrainingFeedback(
        user_id=user.id,
        profile_id=profile.id,
        plan_id=plan_id,
        created_at=now_str(),
        session_index=payload['session_index'],
        completion_status=payload['completion_status'],
        completion_ratio=float(ratio),
        fatigue_score=payload['fatigue_score'],
        discomfort_feedback=payload['discomfort_feedback'],
        sleep_status=payload['sleep_status'],
        screenshot_storage_key=screenshot_key,
        screenshot_url=screenshot_url,
        screenshot_metrics_json=json_dump(metrics),
        feedback_note=payload['feedback_note'],
        result_json=json_dump(result),
    )
    db.session.add(feedback_row)
    db.session.commit()
    return jsonify({'feedback_id': feedback_row.id, **result, 'progress': summarize_progress(user.id)['overview']})


@app.route('/api/membership/dev-upgrade', methods=['POST'])
def membership_dev_upgrade():
    user = require_user()
    if app.config['APP_ENV'] == 'production' or not app.config['ALLOW_DEV_VIP']:
        return jsonify({'error': '生产环境请使用真实支付。'}), 403
    data = request.get_json(force=True) if request.is_json else {}
    tier = data.get('tier', 'vip')
    if tier not in {'free', 'vip'}:
        return jsonify({'error': '无效会员等级。'}), 400
    user.membership_tier = tier
    user.updated_at = now_str()
    db.session.commit()
    return jsonify({'ok': True, 'membership_tier': tier})


@app.route('/api/me/entitlements')
def me_entitlements():
    user = require_user()
    return jsonify({
        'tier': user.membership_tier or 'free',
        'features': compute_entitlements(user)
    })


@app.route('/api/billing/web/checkout-session', methods=['POST'])
def create_checkout_session():
    if not stripe_ready():
        return jsonify({'error': '未配置 Stripe，当前环境无法发起真实支付。'}), 400
    user = require_user()
    if not user.email:
        return jsonify({'error': '请先注册账号并绑定邮箱，再进行会员支付。'}), 400
    configure_stripe()
    data = request.get_json(force=True)
    plan = data.get('plan', 'monthly')
    price_id = app.config['STRIPE_PRICE_MONTHLY'] if plan == 'monthly' else app.config['STRIPE_PRICE_YEARLY']
    if not price_id:
        return jsonify({'error': '未配置对应的价格 ID。'}), 400

    customer_id = user.stripe_customer_id
    if not customer_id:
        customer = stripe.Customer.create(email=user.email, name=user.nickname or None, metadata={'user_id': str(user.id)})
        customer_id = customer.id
        user.stripe_customer_id = customer_id
        user.updated_at = now_str()
        db.session.commit()

    checkout = stripe.checkout.Session.create(
        mode='subscription',
        customer=customer_id,
        line_items=[{'price': price_id, 'quantity': 1}],
        metadata={'user_id': str(user.id)},
        client_reference_id=str(user.id),
        success_url=f"{app.config['BASE_URL']}/?billing=success&session_id={{CHECKOUT_SESSION_ID}}",
        cancel_url=f"{app.config['BASE_URL']}/?billing=cancelled",
    )
    return jsonify({'url': checkout.url})


@app.route('/api/billing/web/customer-portal', methods=['POST'])
def create_customer_portal_session():
    if not stripe_ready():
        return jsonify({'error': '未配置 Stripe。'}), 400
    user = require_user()
    if not user.stripe_customer_id:
        return jsonify({'error': '当前账号没有可管理的订阅。'}), 400
    configure_stripe()
    portal = stripe.billing_portal.Session.create(
        customer=user.stripe_customer_id,
        return_url=f"{app.config['BASE_URL']}/?billing=portal_return",
    )
    return jsonify({'url': portal.url})


@app.route('/api/billing/web/webhook', methods=['POST'])
def stripe_webhook():
    if not stripe_ready():
        return jsonify({'error': 'Stripe 未配置。'}), 400
    configure_stripe()
    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature', '')
    secret = app.config['STRIPE_WEBHOOK_SECRET']
    if not secret:
        return jsonify({'error': '未配置 STRIPE_WEBHOOK_SECRET。'}), 400
    try:
        event = stripe.Webhook.construct_event(payload=payload, sig_header=sig_header, secret=secret)
    except Exception as exc:
        return jsonify({'error': f'Webhook 校验失败：{exc}'}), 400
    process_stripe_event(event)
    return jsonify({'received': True})


with app.app_context():
    db.create_all()


if __name__ == '__main__':
    port = int(os.getenv('PORT', 5050))
    debug = app.config['APP_ENV'] != 'production'
    app.run(debug=debug, host='127.0.0.1', port=port, use_reloader=False)
